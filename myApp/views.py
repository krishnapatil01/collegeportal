from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from .models import RegistrationRequest
from .forms import RegistrationForm, LoginForm, AdminLoginForm
from django.contrib.auth.decorators import login_required

# ===== Resume Analyzer and Chatbot additions =====
import os
import re
from django.conf import settings
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt



def index(request):
    return render(request, 'index.html')

def about(request):
    return render(request, 'about.html')

# Registration request view
def register(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            form.save()  # Save the request for admin approval
            return render(request, 'registration_pending.html')
    else:
        form = RegistrationForm()
    return render(request, 'register.html', {'form': form})


# Admin login view
def admin_login(request):
    if request.method == 'POST':
        form = AdminLoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, username=username, password=password)
            if user and user.is_superuser:
                login(request, user)
                return redirect('admin_dashboard')
            else:
                return render(request, 'Admin/admin_login.html', {'form': form, 'error': 'Invalid credentials or not an admin user.'})
    else:
        form = AdminLoginForm()
    return render(request, 'Admin/admin_login.html', {'form': form})

# Admin dashboard to manage requests
@login_required
def admin_dashboard(request):
    if not request.user.is_superuser:
        return redirect('login')

    requests = RegistrationRequest.objects.filter(is_approved=False)
    return render(request, 'Admin/admin_dashboard.html', {'requests': requests})

# Approve or reject a request
@login_required
def manage_request(request, request_id, action):
    if not request.user.is_superuser:
        return redirect('login')

    req = RegistrationRequest.objects.get(id=request_id)
    if action == 'approve':
        # Create a new user on approval
        User.objects.create_user(
            username=req.username,
            password=req.password,
            email=req.email
        )
        req.is_approved = True
        req.save()
    elif action == 'reject':
        req.delete()

    return redirect('admin_dashboard')
from django.contrib.admin.views.decorators import staff_member_required

@staff_member_required
def view_student_profiles(request):
    """View for displaying and managing student profiles."""
    query = request.GET.get('search', '')  # Get the search query from the URL
    if query:
        students = StudentProfile.objects.filter(
            full_name__icontains=query
        ) | StudentProfile.objects.filter(
            department__icontains=query
        )
    else:
        students = StudentProfile.objects.all()

    if request.method == 'POST':
        # Handle deletion
        student_id = request.POST.get('student_id')
        if student_id:
            student = get_object_or_404(StudentProfile, id=student_id)
            student.user.delete()  # Delete the associated user
            student.delete()       # Delete the profile
            return redirect('view_student_profiles')

    return render(request, 'Admin/admin_view_student.html', {'students': students, 'query': query})
   # return render(request, 'admin_view_student.html', {'students': students})

@staff_member_required
def view_alumni_profiles(request):
    """View for displaying and managing alumni profiles."""
    query = request.GET.get('search', '')  # Get the search query from the URL
    if query:
        alumni = AlumniProfile.objects.filter(
            full_name__icontains=query
        ) | AlumniProfile.objects.filter(
            company_name__icontains=query
        )
    else:
        alumni = AlumniProfile.objects.all()

    if request.method == 'POST':
        # Handle deletion
        alumni_id = request.POST.get('alumni_id')
        if alumni_id:
            alum = get_object_or_404(AlumniProfile, id=alumni_id)
            alum.user.delete()  # Delete the associated user
            alum.delete()       # Delete the profile
            return redirect('view_alumni_profiles')

    return render(request, 'Admin/admin_view_alumni.html', {'alumni': alumni, 'query': query})

from django.contrib import messages
from .models import FundingInfo

@staff_member_required
def add_funding_info(request):
    """Admin view to add funding information."""
    if request.method == 'POST':
        title = request.POST.get('title')
        description = request.POST.get('description')
        amount_required = request.POST.get('amount_required')
        qr_code = request.FILES.get('qr_code')

        funding = FundingInfo(
            title=title,
            description=description,
            amount_required=amount_required,
            qr_code=qr_code,
        )
        funding.save()
        messages.success(request, "Funding information added successfully!")
        return redirect('add_funding_info')

    return render(request, 'Admin/add_funding_info.html')


# Login view
def user_login(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, username=username, password=password)
            if user:
                login(request, user)

                # Fetch the user's registration details
                try:
                    registration_request = RegistrationRequest.objects.get(username=user.username)
                except RegistrationRequest.DoesNotExist:
                    return render(request, 'login.html', {
                        'form': form, 
                        'error': 'Your account is not associated with a valid registration request.'
                    })

                # Redirect based on role and approval status
                if registration_request.is_approved:
                    if registration_request.role == 'student':
                        return redirect('student_dashboard')
                    elif registration_request.role == 'alumni':
                        return redirect('alumni_dashboard')
                else:
                    return render(request, 'login.html', {
                        'form': form, 
                        'error': 'Your registration is pending approval.'
                    })
            else:
                return render(request, 'login.html', {'form': form, 'error': 'Invalid credentials.'})
    else:
        form = LoginForm()
    return render(request, 'login.html', {'form': form})



# Logout view
def user_logout(request):
    logout(request)
    return redirect('login')

# Student dashboard
@login_required
def student_dashboard(request):
    return render(request, 'student/student_dashboard.html')

from .forms import StudentProfileForm
from .models import StudentProfile

@login_required
def create_student_profile(request):
    try:
        student_profile = StudentProfile.objects.get(user=request.user)
        return redirect('view_student_profile')  # Redirect if the profile already exists
    except StudentProfile.DoesNotExist:
        pass

    if request.method == 'POST':
        form = StudentProfileForm(request.POST)
        if form.is_valid():
            profile = form.save(commit=False)
            profile.user = request.user  # Set the user to the current logged-in user
            profile.save()
            return redirect('view_student_profile')
    else:
        form = StudentProfileForm()

    return render(request, 'student/create_student_profile.html', {'form': form})

@login_required
def view_student_profile(request):
    try:
        student_profile = StudentProfile.objects.get(user=request.user)
    except StudentProfile.DoesNotExist:
        return redirect('create_student_profile')  # Redirect if no profile exists

    return render(request, 'student/view_student_profile.html', {'profile': student_profile})

@login_required
def update_student_profile(request):
    profile, created = StudentProfile.objects.get_or_create(user=request.user)
    if request.method == 'POST':
        form = StudentProfileForm(request.POST, request.FILES, instance=profile)  # Include request.FILES to handle image uploads
        if form.is_valid():
            form.save()
            return redirect('view_student_profile')  # Redirect to view the profile after saving
    else:
        form = StudentProfileForm(instance=profile)
    return render(request, 'student/update_student_profile.html', {'form': form})


# View for viewing a specific post and its comments
@login_required
def student_job_port(request):
    posts = Post.objects.all().order_by('-post_date')
    return render(request, 'student/student_job_port.html', {'posts': posts})

@login_required
def view_post(request, post_id):
    from .models import Post, Comment
    from .forms import CommentForm

    post = Post.objects.get(id=post_id)
    comments = Comment.objects.filter(post=post).order_by('-created_at')

    if request.method == 'POST':
        comment_form = CommentForm(request.POST)
        if comment_form.is_valid():
            comment = comment_form.save(commit=False)
            comment.user = request.user  # Assign the logged-in user
            comment.post = post  # Associate the comment with the current post
            comment.save()  # Save the comment
            return redirect('view_post', post_id=post.id)
    else:
        comment_form = CommentForm()

    return render(request, 'student/view_post.html', {
        'post': post,
        'comments': comments,
        'comment_form': comment_form,
    })


def search_alumni(request):
    query = request.GET.get('q', '')  # Get the search query
    alumni_profiles = AlumniProfile.objects.all()  # Default to all profiles

    if query:
        alumni_profiles = alumni_profiles.filter(
            company_name__icontains=query  # Search by company name
        ) | alumni_profiles.filter(
            location__icontains=query  # Search by location
        )

    return render(request, 'student/search_alumni.html', {
        'alumni_profiles': alumni_profiles,
        'query': query,
    })


from django.shortcuts import render, get_object_or_404
from .models import AlumniProfile, Post

def alumni_profile(request, alumni_id):
    # Retrieve the AlumniProfile
    alumni = get_object_or_404(AlumniProfile, id=alumni_id)
    
    # Retrieve posts by matching the user field
    alumni_posts = Post.objects.filter(user=alumni.user)  # Get all posts of the same user as the alumni

    return render(request, 'student/alumni_profile_detail.html', {
        'alumni': alumni,
        'posts': alumni_posts
    })


# Student: View resources
def view_resources(request):
    resources = Resource.objects.all()
    for resource in resources:
        if resource.file.name.lower().endswith(".pdf"):
            resource.file_type = "pdf"
        elif resource.file.name.lower().endswith((".jpg", ".jpeg", ".png", ".gif")):
            resource.file_type = "image"
        elif resource.file.name.lower().endswith((".mp4", ".webm", ".ogg")):
            resource.file_type = "video"
        else:
            resource.file_type = "other"
    return render(request, 'student/view_resources.html', {'resources': resources})

# Alumni dashboard
@login_required
def alumni_dashboard(request):
    return render(request, 'alumni/alumni_dashboard.html')



from django.contrib.auth.decorators import login_required
from .forms import AlumniProfileForm
from .models import AlumniProfile

@login_required
def create_profile(request):
    # Check if the user already has a profile
    if AlumniProfile.objects.filter(user=request.user).exists():
        return redirect('view_alumni_profile')  # Redirect to profile view if profile already exists

    if request.method == 'POST':
        form = AlumniProfileForm(request.POST, request.FILES)  # Include request.FILES to handle file upload
        if form.is_valid():
            profile = form.save(commit=False)
            profile.user = request.user  # Associate the profile with the logged-in user
            profile.save()
            return redirect('view_alumni_profile')  # Redirect to the profile view after saving
    else:
        form = AlumniProfileForm()

    return render(request, 'alumni/create_profile.html', {'form': form})

@login_required
def view_alumni_profile(request):
    try:
        profile = AlumniProfile.objects.get(user=request.user)
    except AlumniProfile.DoesNotExist:
        # If no profile exists, redirect to create profile
        return redirect('create_profile')  # You can implement a create_profile view if necessary
    return render(request, 'alumni/alumni_profile.html', {'profile': profile})

@login_required
def update_alumni_profile(request):
    profile, created = AlumniProfile.objects.get_or_create(user=request.user)
    if request.method == 'POST':
        form = AlumniProfileForm(request.POST, request.FILES, instance=profile)  # Include request.FILES to handle image uploads
        if form.is_valid():
            form.save()
            return redirect('view_alumni_profile')  # Redirect to view the profile after saving
    else:
        form = AlumniProfileForm(instance=profile)
    return render(request, 'alumni/update_alumni_profile.html', {'form': form})

from .forms import PostForm, CommentForm
from .models import Post, Comment

@login_required
def create_post(request):
    if request.method == 'POST':
        form = PostForm(request.POST)
        if form.is_valid():
            post = form.save(commit=False)
            post.user = request.user
            post.save()
            return redirect('post_board')
    else:
        form = PostForm()
    return render(request, 'alumni/create_post.html', {'form': form})

def post_board(request):
    posts = Post.objects.all().order_by('-post_date')
    return render(request, 'alumni/post_board.html', {'posts': posts})

from django.shortcuts import get_object_or_404

@login_required
def edit_post(request, post_id):
    post = get_object_or_404(Post, id=post_id, user=request.user)
    if request.method == 'POST':
        form = PostForm(request.POST, instance=post)
        if form.is_valid():
            form.save()
            return redirect('post_board')
    else:
        form = PostForm(instance=post)
    return render(request, 'alumni/edit_post.html', {'form': form})

@login_required
def view_post_alumini(request, post_id):
    from .models import Post, Comment
    from .forms import CommentForm

    post = Post.objects.get(id=post_id)
    comments = Comment.objects.filter(post=post).order_by('-created_at')

    if request.method == 'POST':
        comment_form = CommentForm(request.POST)
        if comment_form.is_valid():
            comment = comment_form.save(commit=False)
            comment.user = request.user  # Assign the logged-in user
            comment.post = post  # Associate the comment with the current post
            comment.save()  # Save the comment
            return redirect('view_post_alumini', post_id=post.id)
    else:
        comment_form = CommentForm()

    return render(request, 'alumni/view_post_alumini.html', {
        'post': post,
        'comments': comments,
        'comment_form': comment_form,
    })


from .models import Resource
from .forms import ResourceForm

# Alumni: Upload resources
@login_required
def upload_resource(request):
    if request.method == 'POST':
        form = ResourceForm(request.POST, request.FILES)
        if form.is_valid():
            resource = form.save(commit=False)
            resource.uploaded_by = request.user
            resource.save()
            return redirect('alumni_dashboard')  # Change this to your alumni dashboard URL
    else:
        form = ResourceForm()
    return render(request, 'alumni/upload_resource.html', {'form': form})


# Alumni: View Uploaded Notes
@login_required
def view_uploaded_notes(request):
   # resources = Resource.objects.filter(uploaded_by=request.user)
    resources = Resource.objects.all().order_by('-uploaded_at')
    for resource in resources:
        if resource.file.name.lower().endswith(".pdf"):
            resource.file_type = "pdf"
        elif resource.file.name.lower().endswith((".jpg", ".jpeg", ".png", ".gif")):
            resource.file_type = "image"
        elif resource.file.name.lower().endswith((".mp4", ".webm", ".ogg")):
            resource.file_type = "video"
        else:
            resource.file_type = "other"
    return render(request, 'alumni/view_uploaded_notes.html', {'resources': resources})

# Alumni: Edit Uploaded Note
@login_required
def edit_resource(request, pk):
    resource = get_object_or_404(Resource, pk=pk, uploaded_by=request.user)
    if request.method == 'POST':
        form = ResourceForm(request.POST, request.FILES, instance=resource)
        if form.is_valid():
            form.save()
            return redirect('view_uploaded_notes')
    else:
        form = ResourceForm(instance=resource)
    return render(request, 'alumni/edit_resource.html', {'form': form})

# Alumni: Delete Uploaded Note
@login_required
def delete_resource(request, pk):
    resource = get_object_or_404(Resource, pk=pk, uploaded_by=request.user)
    if request.method == 'POST':
        resource.delete()
        return redirect('view_uploaded_notes')
    return render(request, 'alumni/delete_resource.html', {'resource': resource})


def view_funding_info(request):
    """View for alumni to see funding information."""
    funding_list = FundingInfo.objects.all()
    return render(request, 'alumni/view_funding_info.html', {'funding_list': funding_list})

# ===== Resume Analyzer and Chatbot additions =====
def _clean_extracted_text(text):
    """Normalize extracted resume text."""
    text = text or ""
    text = text.replace('\x00', ' ')
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def _extract_text_from_pdf(file_path):
    """
    Extract text from PDF.
    Priority:
    1) PyPDF2 for normal text-based PDFs
    2) PyMuPDF for better PDF text extraction
    3) OCR fallback using PyMuPDF image rendering + pytesseract for scanned/image PDFs
    """
    text = ""

    # 1) Normal text-based PDF extraction
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
    except Exception:
        text = ""

    text = _clean_extracted_text(text)
    if len(text) >= 40:
        return text

    # 2) Better PDF extraction using PyMuPDF if available
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        pymupdf_text = []
        for page in doc:
            pymupdf_text.append(page.get_text("text"))
        doc.close()
        text = _clean_extracted_text("\n".join(pymupdf_text))
        if len(text) >= 40:
            return text
    except Exception:
        pass

    # 3) OCR fallback for scanned PDF/images inside PDF
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
        import io

        doc = fitz.open(file_path)
        ocr_pages = []
        for page in doc:
            # 2x zoom gives better OCR readability
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            page_text = pytesseract.image_to_string(image, lang='eng')
            ocr_pages.append(page_text)
        doc.close()
        text = _clean_extracted_text("\n".join(ocr_pages))
        if len(text) >= 20:
            return text
    except Exception:
        pass

    return ""


def _extract_text_from_docx(file_path):
    """Extract text from DOCX file including paragraphs and tables."""
    try:
        from docx import Document
        document = Document(file_path)
        parts = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    parts.append(' | '.join(row_text))
        return _clean_extracted_text("\n".join(parts))
    except Exception:
        return ""


def _extract_resume_text(uploaded_file, saved_path):
    """Extract resume text from PDF, DOCX, or TXT."""
    file_name = uploaded_file.name.lower()
    if file_name.endswith('.pdf'):
        return _extract_text_from_pdf(saved_path)
    if file_name.endswith('.docx'):
        return _extract_text_from_docx(saved_path)
    if file_name.endswith('.txt'):
        try:
            with open(saved_path, 'r', encoding='utf-8', errors='ignore') as f:
                return _clean_extracted_text(f.read())
        except Exception:
            return ""
    return ""


def _extract_email(text):
    """Extract first email address from resume text."""
    match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text or '')
    return match.group(0) if match else ''


def _extract_phone(text):
    """Extract likely Indian/global phone number from resume text."""
    match = re.search(r'(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{3,5}\)?[\s-]?)?\d{3,5}[\s-]?\d{4,6}', text or '')
    return match.group(0).strip() if match else ''


def _keyword_set(text):
    """Create clean keyword set from text for parser and job matching."""
    words = re.findall(r"[a-zA-Z][a-zA-Z0-9+#.\-]{1,}", (text or "").lower())
    stop_words = {
        'and', 'the', 'for', 'with', 'from', 'this', 'that', 'your', 'you', 'are', 'was', 'were',
        'have', 'has', 'had', 'will', 'can', 'our', 'job', 'role', 'work', 'year', 'years', 'skills',
        'experience', 'responsibility', 'responsibilities', 'description', 'qualification', 'location',
        'email', 'mobile', 'phone', 'address', 'resume', 'curriculum', 'vitae', 'career', 'objective',
        'college', 'university', 'student', 'project', 'projects', 'company', 'name', 'contact', 'profile',
        'summary', 'declaration', 'date', 'place', 'page', 'www', 'com'
    }
    return {w for w in words if len(w) > 2 and w not in stop_words}


def _extract_section(text, section_names, stop_sections=None, max_chars=900):
    """Extract rough resume section text between common headings."""
    if not text:
        return ''
    stop_sections = stop_sections or [
        'education', 'academic', 'qualification', 'skills', 'technical skills', 'experience', 'internship',
        'project', 'projects', 'certification', 'certifications', 'achievement', 'achievements',
        'personal details', 'declaration', 'hobbies', 'languages', 'objective', 'summary'
    ]
    pattern = r'(?i)(?:^|\n|\r|\.|\s)(' + '|'.join(re.escape(x) for x in section_names) + r')\s*[:\-]?\s*'
    match = re.search(pattern, text)
    if not match:
        return ''
    start = match.end()
    remaining = text[start:start + max_chars]
    stop_pattern = r'(?i)(?:\n|\r|\.|\s{2,})(' + '|'.join(re.escape(x) for x in stop_sections if x.lower() not in [n.lower() for n in section_names]) + r')\s*[:\-]?'
    stop_match = re.search(stop_pattern, remaining)
    if stop_match:
        remaining = remaining[:stop_match.start()]
    return _clean_extracted_text(remaining)


def _parse_resume(resume_text):
    """
    Rule-based resume parser technique:
    1) Extract contact details.
    2) Detect technical skills from a controlled skill dictionary.
    3) Extract education, experience, project and certification sections.
    4) Create a normalized keyword set for job matching.
    """
    text = _clean_extracted_text(resume_text)
    lower_text = text.lower()

    skill_dictionary = [
        'python', 'java', 'c', 'c++', 'c#', 'django', 'flask', 'fastapi', 'html', 'css', 'javascript',
        'typescript', 'react', 'angular', 'node', 'express', 'php', 'laravel', 'sql', 'mysql', 'postgresql',
        'mongodb', 'oracle', 'sqlite', 'machine learning', 'deep learning', 'artificial intelligence', 'ai',
        'ml', 'data science', 'data analytics', 'pandas', 'numpy', 'matplotlib', 'tensorflow', 'keras',
        'pytorch', 'opencv', 'nlp', 'computer vision', 'power bi', 'powerbi', 'tableau', 'excel',
        'aws', 'azure', 'google cloud', 'gcp', 'git', 'github', 'docker', 'kubernetes', 'linux',
        'android', 'kotlin', 'xml', 'firebase', 'rest api', 'api', 'bootstrap', 'tailwind', 'jquery',
        'selenium', 'testing', 'manual testing', 'automation testing', 'cyber security', 'networking',
        'communication', 'leadership', 'problem solving', 'teamwork'
    ]

    found_skills = []
    for skill in skill_dictionary:
        skill_pattern = r'(?<![a-z0-9+#])' + re.escape(skill.lower()) + r'(?![a-z0-9+#])'
        if re.search(skill_pattern, lower_text):
            found_skills.append(skill)

    education_keywords = ['b.e', 'be ', 'btech', 'b.tech', 'm.e', 'mtech', 'm.tech', 'bsc', 'b.sc', 'msc', 'm.sc', 'bca', 'mca', 'diploma', 'degree', 'hsc', 'ssc', 'graduation']
    education_matches = [edu.upper() for edu in education_keywords if edu.strip().lower() in lower_text]

    exp_match = re.search(r'(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)', lower_text)
    experience_years = exp_match.group(0) if exp_match else ''

    parsed = {
        'email': _extract_email(text),
        'phone': _extract_phone(text),
        'skills': sorted(set(found_skills)),
        'education_matches': sorted(set(education_matches)),
        'experience_years': experience_years,
        'summary': _extract_section(text, ['summary', 'career objective', 'objective', 'profile summary'], max_chars=600),
        'education_section': _extract_section(text, ['education', 'academic qualification', 'qualification'], max_chars=900),
        'experience_section': _extract_section(text, ['experience', 'work experience', 'internship', 'professional experience'], max_chars=900),
        'project_section': _extract_section(text, ['project', 'projects', 'academic projects'], max_chars=900),
        'certification_section': _extract_section(text, ['certification', 'certifications', 'courses'], max_chars=700),
        'keywords': sorted(_keyword_set(text)),
    }
    return parsed


def _recommend_jobs_from_resume(resume_text, parsed_resume=None):
    """Recommend jobs using parsed resume skills + keywords against existing Post records."""
    parsed_resume = parsed_resume or _parse_resume(resume_text)
    resume_keywords = set(parsed_resume.get('keywords', []))
    parsed_skills = {skill.lower() for skill in parsed_resume.get('skills', [])}

    # Add split skill words also, so "machine learning" can match either exact phrase or job text words.
    for skill in list(parsed_skills):
        resume_keywords.update(_keyword_set(skill))

    recommendations = []
    for post in Post.objects.all().order_by('-post_date'):
        job_text = " ".join([
            post.title or "",
            post.company_name or "",
            post.qualification or "",
            post.role_responsibility or "",
            post.skills or "",
            post.description or "",
        ])
        job_lower = job_text.lower()
        job_keywords = _keyword_set(job_text)

        exact_skill_matches = sorted([skill for skill in parsed_skills if skill and skill in job_lower])
        keyword_matches = sorted(resume_keywords.intersection(job_keywords))

        # Parser-based scoring: exact parsed skills are more important than generic words.
        score = (len(exact_skill_matches) * 5) + len(keyword_matches)
        matched_display = sorted(set(exact_skill_matches + keyword_matches))[:20]

        if score > 0:
            recommendations.append({
                'post': post,
                'score': score,
                'matched_skills': matched_display,
            })

    recommendations.sort(key=lambda item: item['score'], reverse=True)
    return recommendations[:10]

def _latest_jobs(limit=6):
    """Fallback jobs to show when no exact resume match is found."""
    latest = []
    for post in Post.objects.all().order_by('-post_date')[:limit]:
        latest.append({
            'post': post,
            'score': 0,
            'matched_skills': ['latest job post'],
        })
    return latest


@login_required
def upload_resume(request):
    """Student resume upload + content extraction + job recommendation."""
    if request.method == 'POST':
        resume_file = request.FILES.get('resume_file')
        if not resume_file:
            return render(request, 'student/upload_resume.html', {'error': 'Please select a resume file.'})

        allowed_extensions = ('.pdf', '.docx', '.txt')
        if not resume_file.name.lower().endswith(allowed_extensions):
            return render(request, 'student/upload_resume.html', {
                'error': 'Please upload only PDF, DOCX or TXT resume file.'
            })

        resume_dir = os.path.join(settings.MEDIA_ROOT, 'resumes')
        os.makedirs(resume_dir, exist_ok=True)
        safe_name = re.sub(r'[^a-zA-Z0-9_.-]', '_', resume_file.name)
        saved_path = os.path.join(resume_dir, f"{request.user.username}_{safe_name}")
        with open(saved_path, 'wb+') as destination:
            for chunk in resume_file.chunks():
                destination.write(chunk)

        resume_text = _extract_resume_text(resume_file, saved_path)
        parsed_resume = _parse_resume(resume_text) if resume_text else {}
        recommendations = _recommend_jobs_from_resume(resume_text, parsed_resume) if resume_text else []
        fallback_jobs = [] if recommendations else _latest_jobs()

        extraction_message = ''
        if not resume_text and resume_file.name.lower().endswith('.pdf'):
            extraction_message = (
                'Resume text could not be extracted automatically. This PDF may be scanned/image-based. '
                'Install Tesseract OCR and keep PyMuPDF + pytesseract installed, or upload DOCX/TXT resume.'
            )
        elif not resume_text:
            extraction_message = 'Resume text could not be extracted. Please upload a text-based PDF, DOCX, or TXT resume.'

        return render(request, 'student/analyze_resume.html', {
            'resume_file': resume_file.name,
            'resume_text': resume_text,
            'parsed_resume': parsed_resume,
            'recommendations': recommendations,
            'fallback_jobs': fallback_jobs,
            'extraction_message': extraction_message,
        })

    return render(request, 'student/upload_resume.html')


def chatbot(request):
    """Simple career chatbot page."""
    return render(request, 'student/chatbot.html')



import os
from django.conf import settings
from django.http import JsonResponse
from google import genai

@csrf_exempt
def chatbot_response(request):
    """Return chatbot answer using Gemini 2.5 Flash for student/alumni portal questions."""

    if request.method == 'POST':
        user_message = request.POST.get('message', '')
    else:
        user_message = request.GET.get('message', '')

    user_message = user_message.strip()

    if not user_message:
        return JsonResponse({
            'response': 'Please type your question about resume upload, jobs, alumni search, notes, login, or profile.'
        })

    # Get Gemini API key from settings.py or environment variable
    api_key = "AIzaSyCY3Vi1WrFx5_7F1PcwCSi89nMKMW6e6YQ"

    if not api_key:
        return JsonResponse({
            'response': 'Gemini API key is missing. Please add GEMINI_API_KEY in settings.py or environment variable.'
        })

    try:
        client = genai.Client(api_key=api_key)

        prompt = f"""
You are a helpful chatbot for a Student Alumni Portal website.

Education guidance details:
- Help students with education-related questions.
- Guide students about courses, skills, projects, internships, resume building, career paths, and job preparation.
- Suggest skills based on education background such as B.Tech, BCA, BCS, MCA, MBA, Diploma, Computer Science, IT, AI/ML, Data Science, Web Development, Cyber Security, Cloud Computing, and Software Testing.
- Explain which technical skills students should learn for job readiness.
- Suggest project ideas for final-year students.
- Suggest internship domains based on student interest.
- Help students understand how alumni can support them with career guidance.
- Help students prepare for interviews, aptitude, technical questions, and communication skills.
- Help students choose suitable job roles based on education and resume skills.
- If user asks about education details from resume, explain education section such as degree, college, university, percentage/CGPA, passing year, and specialization.
- Keep education answers simple, practical, and student-friendly.

Rules:
- Answer only questions related to this portal, education guidance, resume, jobs, internships, alumni, notes, profile, or login.
- Keep answers short, simple, and student-friendly.
- If user asks about jobs, explain job board and resume-based recommendations.
- If user asks about education, give useful course, skill, project, internship, and career guidance.
- If user asks about alumni, explain alumni search and connection.
- If user asks unrelated questions, politely say you can help only with student/alumni portal and education/career guidance.

User question:
{user_message}
"""

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )

        answer = response.text.strip() if response.text else "Sorry, I could not generate a response."

    except Exception as e:
        print("Gemini Chatbot Error:", e)

        # Backup response if Gemini API fails
        msg = user_message.lower()

        if any(word in msg for word in ['resume', 'cv', 'upload']):
            answer = 'Go to Student Dashboard and click Upload Resume. Upload a PDF, DOCX, or TXT resume.'
        elif any(word in msg for word in ['job', 'internship', 'recommend', 'opening']):
            answer = 'Use Job & Internship Board to view posts. For personalized recommendations, upload your resume.'
        elif any(word in msg for word in ['profile', 'update profile']):
            answer = 'Open Student Dashboard and click Update Profile to edit your student details.'
        elif any(word in msg for word in ['alumni', 'search', 'connect']):
            answer = 'Click Search Alumni from the student dashboard to find alumni by company or location.'
        elif any(word in msg for word in ['notes', 'resource', 'study']):
            answer = 'Click Notes in the student navbar to view resources uploaded by alumni.'
        elif any(word in msg for word in ['login', 'register', 'account']):
            answer = 'Register as Student or Alumni first. After admin approval, login using username and password.'
        else:
            answer = 'I can help with resume upload, job recommendations, student profile, alumni search, notes, login, and portal navigation.'

    return JsonResponse({'response': answer})